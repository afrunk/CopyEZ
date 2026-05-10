"""
Quotes (摘抄语录) module.

Displays all highlighted quotes with annotations, grouped by article.
Supports export to Word documents.
"""

import re
import os
from flask import Blueprint, render_template, make_response

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.shared import Pt
except ImportError:
    Document = None
    WD_ALIGN_PARAGRAPH = None
    WD_LINE_SPACING = None
    qn = None
    Pt = None

from app.modules.copyez.quote import collect_quote_items

quotes_bp = Blueprint("quotes", __name__)


@quotes_bp.route("/my_quotes")
def my_quotes_page():
    """
    摘抄语录本：集中展示所有带批注的高亮句子，按"文章"分组。
    """
    article_groups = collect_quote_items()
    return render_template("my_quotes.html", article_groups=article_groups)


@quotes_bp.route("/my_quotes/export")
def export_my_quotes():
    """
    将所有摘抄导出为 Word 文档（公文友好格式），按"文章"分组。
    """
    article_groups = collect_quote_items()

    if not Document:
        return (
            "当前环境未安装 python-docx，请先在服务器上执行 pip install python-docx 后重试导出。",
            500,
        )

    doc = Document()

    def _set_run_font(run, east_asia_name: str, latin_name: str = None, size_pt: int = None, bold: bool = None):
        if bold is not None:
            run.bold = bool(bold)
        if size_pt and Pt:
            run.font.size = Pt(size_pt)
        if latin_name:
            run.font.name = latin_name
        if qn:
            r = run._element
            rPr = r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn("w:eastAsia"), east_asia_name)

    def _set_paragraph_format(p, line_spacing_pt: int = 28):
        if not p:
            return
        try:
            if WD_LINE_SPACING:
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            if Pt:
                p.paragraph_format.line_spacing = Pt(line_spacing_pt)
        except Exception:
            pass

    try:
        normal = doc.styles["Normal"]
        if Pt:
            normal.font.size = Pt(12)
        normal.font.name = "FangSong"
        if qn:
            normal._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
    except Exception:
        pass

    title_p = doc.add_paragraph()
    run = title_p.add_run("摘抄语录本")
    _set_run_font(run, east_asia_name="黑体", latin_name="SimHei", size_pt=16, bold=True)
    if WD_ALIGN_PARAGRAPH:
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_format(title_p, 28)

    for group in article_groups:
        heading_p = doc.add_paragraph()
        heading_run = heading_p.add_run(str(group.get("note_title") or "未命名文章"))
        _set_run_font(heading_run, east_asia_name="黑体", latin_name="SimHei", size_pt=14, bold=True)
        _set_paragraph_format(heading_p, 28)

        items = group.get("items") or []
        for idx, q in enumerate(items, start=1):
            p = doc.add_paragraph()
            run_text_label = p.add_run(f"{idx}. 原文：")
            _set_run_font(run_text_label, east_asia_name="黑体", latin_name="SimHei", bold=True)
            quote_plain = re.sub(r"<[^>]+>", "", q.get("text") or "")
            run_text = p.add_run(quote_plain)
            _set_run_font(run_text, east_asia_name="仿宋", latin_name="FangSong")
            p.add_run("\n")
            if q.get("comment"):
                run_c_label = p.add_run("   批注：")
                _set_run_font(run_c_label, east_asia_name="黑体", latin_name="SimHei", bold=True)
                run_c = p.add_run(q["comment"])
                _set_run_font(run_c, east_asia_name="仿宋", latin_name="FangSong")
                p.add_run("\n")
            run_src_label = p.add_run("   来源：")
            _set_run_font(run_src_label, east_asia_name="黑体", latin_name="SimHei", bold=True)
            run_src = p.add_run(f"《{q['note_title']}》")
            _set_run_font(run_src, east_asia_name="仿宋", latin_name="FangSong")
            _set_paragraph_format(p, 28)

    output_path = os.path.join(os.getcwd(), "my_quotes_export.docx")
    doc.save(output_path)

    with open(output_path, "rb") as f:
        data = f.read()

    response = make_response(data)
    response.headers["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    response.headers["Content-Disposition"] = "attachment; filename=my_quotes_export.docx"
    return response
